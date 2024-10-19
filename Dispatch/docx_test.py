from docx import Document
from docx.shared import Inches

doc = Document()
doc.add_heading(text='455', level=0)
doc.add_heading(text='455', level=1)
doc.add_heading(text='455', level=2)
doc.add_heading(text='455', level=2)
doc.add_paragraph('123')
doc.save('C:/Users/24253/Desktop/Python_World/doc1.docx')