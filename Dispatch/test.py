from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def center_insert_img(doc, img):
    """插入图片"""
    for paragraph in doc.paragraphs:
        # 根据文档中的占位符定位图片插入的位置
        if '<<img1>>' in paragraph.text:
            # 把占位符去掉
            paragraph.text = paragraph.text.replace('<<img1>>', '')
            # 添加一个文字块
            run = paragraph.add_run('')
            # 添加一个’回车换行效果‘
            run.add_break()
            # 添加图片并指定大小
            run.add_picture(img, width=Inches(6.2))


def save_img_to_doc(img):
    """把图片保存到doc文件中的指定位置"""
    tpl_doc = 'C:/Users/24253/Desktop/每月例行工作/10月工作内容/10月派单回单/派单/LTE;城关镇解放军第九五五医院附近;LTE质差;2024年10月 - 副本.docx'
    res_doc = 'C:/Users/24253/Desktop/每月例行工作/10月工作内容/10月派单回单/派单/LTE;城关镇解放军第九五五医院附近;LTE质差;2024年10月 - 副本1.docx'
    # 打开模板文件
    document = Document(tpl_doc)
    # 插入图片居中
    center_insert_img(document, img)
    # 保存结果文件
    document.save(res_doc)


def main():
    """主函数"""
    img = 'C:/Users/24253/Desktop/每月例行工作/10月工作内容/10月派单回单/image/城关镇野猪坝大桥附近道路质差.png'
    save_img_to_doc(img)


if __name__ == '__main__':
    main()



