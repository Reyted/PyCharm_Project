import time
import xlwings as xw
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def creat_report(url_attach,url_tpl,url_img,city):
    wb = xw.Book(
        f'{url_attach}{city}LTE&NR测试分析优化问题点10月.xlsx')
    sht = wb.sheets['Sheet1']
    data_list = sht.range(2, 3).options(expand='table').value
    title_lst=[]
    ind=0
    tpl = DocxTemplate(
        f'{url_tpl}NR;昌都市城关镇北大桥路段;NR弱覆盖;2024年10月_模板.docx')
    for data in data_list:
        pos = data[3].split('道路')[0]
        type = '弱覆盖' if data[8] == '覆盖问题' else '质差'
        datetime = data[0].strftime('%Y年%m月')
        context = {
            'system': data[2],
            'position': pos,
            'question_type': type,
            'length': int(data[6]),
            'datetime': data_list[0][0].strftime('%Y年%m月'),
            'lontitude': data[4],
            'latitude': data[5],
            'serving_cell': data[7],
            'type_define': data[11],
            'problem_analysis': data[9],
            'solution_measures': data[10]
        }

        title_lst=[data[2],pos,type,data_list[0][0].strftime('%Y年%m月'),data_list[ind][3]]


        tpl.render(context)
        tpl.save(
            f'{url_tpl}{data[2]};{pos};{data[2]}{type};{datetime}.docx')
        main(url_img,url_tpl,title_lst,city)
        ind+=1

    wb.close()

def center_insert_img(doc, img):
    """插入图片"""
    for paragraph in doc.paragraphs:
        # 根据文档中的占位符定位图片插入的位置
        if '<<img1>>' in paragraph.text:
            # 把占位符去掉
            paragraph.text = paragraph.text.replace('<<img1>>', '')
            # 添加一个文字块
            run = paragraph.add_run('')
            # 添加一个’回车换行效果‘
            run.add_break()
            # 添加图片并指定大小
            run.add_picture(img, width=Inches(6.2))


def save_img_to_doc(img,url_tpl,title_lst:list):
    """把图片保存到doc文件中的指定位置"""
    tpl_doc = f'{url_tpl}{title_lst[0]};{title_lst[1]};{title_lst[0]}{title_lst[2]};{title_lst[3]}.docx'
    res_doc = f'{url_tpl}{title_lst[0]};{title_lst[1]};{title_lst[0]}{title_lst[2]};{title_lst[3]}.docx'
    # 打开模板文件
    document = Document(tpl_doc)
    # 插入图片居中
    center_insert_img(document, img)
    # 保存结果文件
    document.save(res_doc)


def main(url_img:str,url_tpl,title_lst:list,city:str):
    """主函数"""
    img = f'{url_img}{city}/{title_lst[4]}.png'
    save_img_to_doc(img,url_tpl,title_lst)

if __name__ == '__main__':
    url_attach='C:/Users/24253/Desktop/每月例行工作/10月工作内容/10月工单报告/10月工单报告附件/'
    url_tpl='C:/Users/24253/Desktop/每月例行工作/10月工作内容/10月派单回单/派单/'
    url_img='C:/Users/24253/Desktop/每月例行工作/10月工作内容/10月派单回单/image/'
    city_lst=['昌都市城区','昌都机场高速','昌都市察雅县','昌都至察雅国道']
    for city in city_lst:
        creat_report(url_attach,url_tpl,url_img,city)