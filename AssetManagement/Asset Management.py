import time
import xlwings as xw
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Test import distance
import ScreecShot
import pandas as pd


def creat_report(url_tpl,lst:list,ind:int):
    tpl = DocxTemplate(
        f'{url_tpl}muban.docx')
    context = {
        'lontitude': lst[0],
        'latitude': lst[1],
        'hometown': lst[2],
    }
    tpl.render(context)

    tpl.save(f'{url_tpl}allfile/{ind}.docx')
    save_img_to_doc(f'C:/Users/24253/Desktop/每月例行工作/123/资产标签/{ind}.jpg','C:/Users/24253/Desktop/每月例行工作/123/模板/',ind)

def center_insert_img(doc, img):
    """插入图片"""
    for paragraph in doc.paragraphs:
        # 根据文档中的占位符定位图片插入的位置
        if '<<img>>' in paragraph.text:
            # 把占位符去掉
            paragraph.text = paragraph.text.replace('<<img>>', '')
            # 添加一个文字块
            run = paragraph.add_run('')
            # 添加图片并指定大小
            run.add_picture(img, width=Inches(6.2))


def save_img_to_doc(img,url_tpl,ind):
    """把图片保存到doc文件中的指定位置"""
    tpl_doc = f'{url_tpl}allfile/{ind}.docx'
    res_doc = f'{url_tpl}allfile/{ind}.docx'
    # 打开模板文件
    document = Document(tpl_doc)
    # 插入图片居中
    center_insert_img(document, img)
    # 保存结果文件
    document.save(res_doc)

def read_txt(file_path, encoding='utf-8'):
    try:
        with open(file_path, 'r', encoding=encoding) as file:
            return file.read()
    except Exception as e:
        print(f"读取文件时出错: {str(e)}")
        return None


def read_excel(file_path, sheet_name=0):
    try:
        # 读取Excel文件
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        return df
    except Exception as e:
        print(f"读取Excel文件时出错: {str(e)}")
        return None


def find_string_index(string_list_2d, target_string, case_sensitive=True):
    """
    获取包含指定字符串的第一个匹配项在二维列表中的索引位置
    :param string_list_2d: 二维字符串列表
    :param target_string: 要查找的目标字符串
    :param case_sensitive: 是否区分大小写，默认True
    :return: 索引元组(行号, 列号)，如果未找到返回None
    """
    # 如果不区分大小写，将目标字符串转换为小写
    print(target_string)
    if not case_sensitive:
        target_string = target_string.lower()

    # 遍历二维列表
    for row_idx, row in enumerate(string_list_2d):
        for col_idx, item in enumerate(row):
            # 转换当前项为字符串并根据大小写设置进行比较
            current_item = str(item)
            if not case_sensitive:
                current_item = current_item.lower()

            # 检查是否包含目标字符串，找到就立即返回
            if target_string in current_item:
                return (row_idx, col_idx)

    # 未找到返回None
    return None

# 使用示例
if __name__ == '__main__':
    url_tpl = 'C:/Users/24253/Desktop/每月例行工作/123/模板/'
    content = read_txt('C:/Users/24253/Desktop/每月例行工作/123/模板/6.txt')
    base_station_5 = read_excel('C:/Users/24253/Desktop/每月例行工作/工参/昌都245G工参总版-1028.xlsx', 'NR').values.tolist()
    base_station_4 = read_excel('C:/Users/24253/Desktop/每月例行工作/工参/昌都245G工参总版-1028.xlsx', 'LTE').values.tolist()
    base_station=base_station_5+base_station_4
    lst=content.split('\n')
    ind=1
    for i in lst:
        for item in base_station:
            # if item[0]=='昌都':
            #     str = item[12].split('_')[1] + item[12].split('_')[2]
            # else:
            #     str = item[0].split('_')[1] + item[0].split('_')[2]
            if i in item:
                if item[0]=='昌都':
                    lst=[]
                    lst.append(item[10])
                    lst.append(item[11])
                    try:
                        cun=item[12].split('镇')[1].split('村')
                        if len(cun)>1:
                            str1 = cun[0] + '村'
                            lst.append(item[1] + item[2] + str1)
                        else:
                            lst.append(item[1] + item[2])
                    except:
                        try:
                            cun=item[12].split('乡')[1].split('村')
                            if len(cun) > 1:
                                str1 = cun[0] + '村'
                                lst.append(item[1] + item[2] + str1)
                            else:
                                lst.append(item[1] + item[2])
                        except:
                            lst.append(item[1] + item[2])
                    creat_report(url_tpl,lst,ind)
                    ScreecShot.main(ind)
                    ind+=1
                    break
                else:
                    lst = []
                    lst.append(item[15])
                    lst.append(item[16])
                    st=item[0].split('_')
                    lst.append(st[1]+st[2])
                    creat_report(url_tpl, lst, ind)
                    ScreecShot.main(ind)
                    ind += 1
                    break

